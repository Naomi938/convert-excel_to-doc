# -*- coding: utf-8 -*-
import subprocess, sys
subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "openpyxl"],
               capture_output=True)

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# ── RTL helpers ───────────────────────────────────────────────────────────────
def make_rtl_para(para):
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:jc')):
        pPr.remove(existing)
    for existing in pPr.findall(qn('w:bidi')):
        pPr.remove(existing)
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def make_rtl_run(run, size_pt, bold=False, color=None, font_name='David'):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    for existing in rPr.findall(qn('w:rtl')):
        rPr.remove(existing)
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'),    font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szCs)

def set_doc_defaults_rtl(doc):
    settings_el = doc.settings.element
    for existing in settings_el.findall(qn('w:bidi')):
        settings_el.remove(existing)
    settings_el.append(OxmlElement('w:bidi'))
    try:
        normal_style = doc.styles['Normal']
        pPr = normal_style.element.get_or_add_pPr()
        for existing in pPr.findall(qn('w:bidi')):
            pPr.remove(existing)
        for existing in pPr.findall(qn('w:jc')):
            pPr.remove(existing)
        b = OxmlElement('w:bidi')
        pPr.insert(0, b)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)
        normal_style.font.name = 'David'
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    except Exception:
        pass

# ── Streamlit UI ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="Excel → Word", page_icon="📄", layout="centered")

st.markdown("""
<style>
    body, .stApp { direction: rtl; }
    h1,h2,h3,p,label { text-align: right; }
</style>
""", unsafe_allow_html=True)

st.title("📄 המרת שאלות ותשובות מ-Excel ל-Word")
st.markdown("העלה קובץ Excel עם שאלות בעמודה **A** ותשובות בעמודה **D**")

uploaded_file = st.file_uploader("בחר קובץ Excel", type=["xlsx", "xls"])

if uploaded_file:
    try:
        df_raw = pd.read_excel(uploaded_file, header=None, dtype=str)
    except Exception as e:
        st.error(f"שגיאה בקריאת הקובץ: {e}")
        st.stop()

    total_rows = len(df_raw)
    st.success(f"✅ הקובץ נטען — {total_rows} שורות")

    # סדר: 12, 5, 3, 4, 6–11, 13+ (ללא שורות 1 ו-2)
    row_12    = df_raw.iloc[[11]]
    row_5     = df_raw.iloc[[4]]
    row_3     = df_raw.iloc[[2]]
    row_4     = df_raw.iloc[[3]]
    rows_6_11 = df_raw.iloc[5:11]
    rows_13on = df_raw.iloc[12:]
    df_ordered = pd.concat([row_12, row_5, row_3, row_4, rows_6_11, rows_13on], ignore_index=True)

    questions = df_ordered.iloc[:, 0]
    answers   = df_ordered.iloc[:, 3] if df_ordered.shape[1] > 3 else pd.Series([""] * len(df_ordered))

    qa_pairs = []
    for q, a in zip(questions, answers):
        q_str = str(q).strip() if pd.notna(q) else ""
        a_str = str(a).strip() if pd.notna(a) else ""
        if q_str and q_str.lower() != "nan":
            qa_pairs.append((q_str, a_str if a_str.lower() != "nan" else ""))

    if not qa_pairs:
        st.warning("לא נמצאו שאלות תקינות בעמודה A")
        st.stop()

    st.markdown("### תצוגה מקדימה (5 ראשונות)")
    st.table([{"שאלה": q, "תשובה": a} for q, a in qa_pairs[:5]])
    st.caption(f'סה"כ {len(qa_pairs)} שאלות ותשובות')

    # שם התלמיד מעמודה D שורה 12
    student_name = ""
    if df_raw.shape[0] > 11 and df_raw.shape[1] > 3:
        val = str(df_raw.iloc[11, 3]).strip()
        if val and val.lower() != "nan":
            student_name = val
    doc_title = st.text_input("כותרת המסמך", value=student_name or "שאלות ותשובות")

    if st.button("🔄 צור קובץ Word", type="primary"):
        doc = Document()
        set_doc_defaults_rtl(doc)

        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        # כותרת
        title_para = doc.add_paragraph()
        make_rtl_para(title_para)
        title_para.paragraph_format.space_after = Pt(16)
        r = title_para.add_run(doc_title)
        make_rtl_run(r, size_pt=22, bold=True, color=RGBColor(0x2E, 0x40, 0x57))
        doc.add_paragraph()

        # שאלות ותשובות
        for i, (q, a) in enumerate(qa_pairs, start=1):
            # שאלה
            q_para = doc.add_paragraph()
            make_rtl_para(q_para)
            q_para.paragraph_format.space_before = Pt(10)
            q_para.paragraph_format.space_after  = Pt(2)
            r_num = q_para.add_run(f"{i}. ")
            make_rtl_run(r_num, size_pt=13, bold=True, color=RGBColor(0x2E, 0x40, 0x57))
            r_q = q_para.add_run(q)
            make_rtl_run(r_q, size_pt=13, bold=True, color=RGBColor(0x2E, 0x40, 0x57))

            # תשובה
            a_para = doc.add_paragraph()
            make_rtl_para(a_para)
            a_para.paragraph_format.space_after = Pt(2)
            r_label = a_para.add_run("תשובה: ")
            make_rtl_run(r_label, size_pt=12, bold=True, color=RGBColor(0x27, 0xAE, 0x60))
            r_a = a_para.add_run(a)
            make_rtl_run(r_a, size_pt=12, color=RGBColor(0x33, 0x33, 0x33))

            # ספירת מילים
            # הסר תבנית "מספר)" מתחילת התשובה לפני הספירה
            import re
            a_clean = re.sub(r'^\d+\)\s*', '', a.strip())
            word_count = len(a_clean.split()) if a_clean.strip() else 0
            wc_para = doc.add_paragraph()
            make_rtl_para(wc_para)
            wc_para.paragraph_format.space_after = Pt(6)
            wc_run = wc_para.add_run(f"מילים: {word_count}")
            make_rtl_run(wc_run, size_pt=9, color=RGBColor(0x99, 0x99, 0x99))
            wc_run.italic = True

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        st.success(f"✅ המסמך נוצר עם {len(qa_pairs)} שאלות ותשובות!")
        st.download_button(
            label="⬇️ הורד קובץ Word",
            data=buf,
            file_name=f"{doc_title}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
